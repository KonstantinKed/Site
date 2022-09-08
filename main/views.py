from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from .models import Inventory, Pics
from .form import CreateItemForm, CreateBrandForm, UpdateItemForm, UploadPicsForm, ApartmentSelectForm
from django.http import HttpResponseRedirect, HttpResponse
from django.db.models import OuterRef, Subquery, Count, Min
from docx import Document
from docx.shared import Inches

# Create your views here.
from .utils import items_calc


def index(request):
    return render(request, 'index.html')

def control(request):
    items = Inventory.objects.all()
    context = {'items': items}
    return render(request, 'sidebar.html', context=context)

@login_required(login_url='/login')
def items(request):
    items = Inventory.objects.all()
    for item in items:
        pics = Pics.objects.filter(invent = item.id)
    if request.method == 'POST':
        apt_form = ApartmentSelectForm(request.POST) # apartment form to select in items page
        if apt_form.is_valid():
            items = Inventory.objects.filter(apt=apt_form.cleaned_data['apt'])  # selecting items according to selected apartment
    else:
        apt_form = ApartmentSelectForm()
    context = {'items': items, 'apt_form': apt_form}
    return render(request, 'items.html', context=context)

@login_required(login_url='/login')
def create_item(request):
    if request.method == 'POST':
        item_form = CreateItemForm(request.POST)
        pic_form = UploadPicsForm(request.POST, request.FILES)
        if item_form.is_valid():
            item = item_form.save()
            uploaded_img = pic_form.save(commit=False)
            for img in request.FILES.getlist('images'):
                Pics.objects.create(img=img, invent=item)
                # pic_form.cleaned_data['img']
                # uploaded_img.img = img
                # # invent = get_object_or_404(Inventory, id=form.cleaned_data['invent_id'])
                # uploaded_img.invent = item
                # uploaded_img.save()
            return HttpResponseRedirect('/items')
    else:
        item_form = CreateItemForm()
        pic_form = UploadPicsForm()
    context = {'item_form': item_form, 'pic_form': pic_form}
    return render(request, 'create_item.html', context)

# Adding new brands to the list of brands
@login_required(login_url='/login')
def create_brand(request):
    if request.method == 'POST':
        brand_form = CreateBrandForm(request.POST)
        if brand_form.is_valid():
            brand = brand_form.save()
            return HttpResponseRedirect('/items')
    else:
        brand_form = CreateBrandForm()
    return render(request, 'create_brand.html', {'brand_form': brand_form})


@login_required(login_url='/login')
def view_item(request, id):
    instance = get_object_or_404(Inventory, id=id)
    pics = Pics.objects.filter(invent = id)
    # for item in instance:
    #     pics = Pics.objects.filter(invent = item.id)
    # if request.method == 'POST':
    #     form = UpdateItemForm(request.POST, instance=instance)
    #     if form.is_valid():
    #         form.save()
    #         return HttpResponseRedirect('/items')
    # else:
    # form = UpdateItemForm(None, instance=instance)
    return render(request, 'view_item.html', {'instance': instance, 'pics': pics})


@login_required(login_url='/login')
def update_item(request, id):
    instance = get_object_or_404(Inventory, id=id)
    if request.method == 'POST':
        form = UpdateItemForm(request.POST, instance=instance)
        if form.is_valid():
            form.save()
            return HttpResponseRedirect('/items')
    else:
        form = UpdateItemForm(None, instance=instance)
        pics = Pics.objects.filter(invent = id)
    return render(request, 'update_item.html', {'form': form, 'pics': pics})


@login_required(login_url='/login')
def delete_item(request, id):
    instance = get_object_or_404(Inventory, id=id)
    instance.delete()
    return HttpResponseRedirect('/items')


@login_required(login_url='/login')
def upload_pics(request):
    if request.method == 'POST':
        form = UploadPicsForm(request.POST, request.FILES)
        print(form.data)
        if form.is_valid():
            uploaded_img = form.save(commit=False)
            uploaded_img.image_data = form.cleaned_data['img'].file.read()
            invent = Inventory.objects.filter(invent_id = form.cleaned_data['invent_id']).first()
            # invent = get_object_or_404(Inventory, invent_id=form.cleaned_data['invent_id'])
            uploaded_img.invent_id = invent
            uploaded_img.save()
            return HttpResponseRedirect('/')
    else:
        form = UploadPicsForm()
    return render(request, 'upload_pics.html', {'form': form})


# get items, check if duplicate, than counts and prepare for printing.
def collect_list(request):
    items = Inventory.objects.all()
    if request.method == 'POST':
        apt_form = ApartmentSelectForm(request.POST)
        if apt_form.is_valid():
            items = Inventory.objects.filter(
                apt=apt_form.cleaned_data['apt'])
    else:
        apt_form = ApartmentSelectForm()
    inventories = items_calc(items) # function code is in utils fyle
    context = {'items': items, 'apt_form': apt_form, 'inventories': inventories}

    return render(request, 'collect_list.html', context=context)

def generate_docx(request):

    items = Inventory.objects.all()
    inventories = items_calc(items)
    document = Document()

    document.add_heading('Акт прийому-передачі Помешкання', 0)

    p = document.add_paragraph('В  Помешканні знаходсяться наступне Майно: ')
    # p.add_run('bold').bold = True
    # p.add_run(' and some ')
    # p.add_run('italic.').italic = True

    # document.add_heading('Heading, level 1', level=1)
    # document.add_paragraph('Intense quote', style='Intense Quote')
    #
    # document.add_paragraph(
    #     'first item in unordered list', style='List Bullet'
    # )
    for v in inventories.values():
        print(v)
        document.add_paragraph(
        str(v['name']) + '-' + str(v['count']), style='List Number'
        )

    # document.add_picture('monty-truth.png', width=Inches(1.25))
    #
    # records = (
    #     (3, '101', 'Spam'),
    #     (7, '422', 'Eggs'),
    #     (4, '631', 'Spam, spam, eggs, and spam')
    # )
    #
    # table = document.add_table(rows=1, cols=3)
    # hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'Qty'
    # hdr_cells[1].text = 'Id'
    # hdr_cells[2].text = 'Desc'
    # for qty, id, desc in records:
    #     row_cells = table.add_row().cells
    #     row_cells[0].text = str(qty)
    #     row_cells[1].text = id
    #     row_cells[2].text = desc

    document.add_page_break()


    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=download.docx'
    document.save(response)

    return response